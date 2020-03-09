from docx import Document
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import os
import time
import logging
from configparser import ConfigParser

"""settings.ini"""
config = ConfigParser()
configList = config.read("settings.ini")
projectDir = os.getcwd()

if configList.__len__() == 0:
    config["Paths"] = {"Path": projectDir}

    with open("settings.ini", "w") as configfile:
        config.write(configfile)

"""Writing the logs"""
loggerError = logging.getLogger("error")
loggerError.setLevel(logging.ERROR)

loggerInfo = logging.getLogger("info")
loggerInfo.setLevel(logging.INFO)

formatter = logging.Formatter("%(asctime)s:%(message)s")

"""if directory Logs doesn't exist"""
logDir = os.path.abspath("Logs")
if not os.path.isdir(logDir):
    os.mkdir(logDir)

infoHandler = logging.FileHandler(os.path.join("Logs", "info.log"))
infoHandler.setLevel(logging.INFO)
infoHandler.setFormatter(formatter)

errorHandler = logging.FileHandler(os.path.join("Logs", "errors.log"))
errorHandler.setLevel(logging.raiseExceptions)
errorHandler.setFormatter(formatter)

loggerError.addHandler(errorHandler)
loggerInfo.addHandler(infoHandler)

def logDecorator(func):
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except BaseException as e:
            loggerError.exception(f"An error has been occurred in function {func.__name__}", exc_info=e)

    return wrapper

# @logDecorator
# def splitWordFile(filePath):
#     pass
# word = Document(filePath)
# rows = word.tables[0].rows
#
# tbl = word.tables[0]._tbl
# tr = rows[13]._tr
# tbl.remove(tr)
#
# for row in rows:
#     pass

# cells = word.tables[0]._cells
# for cell in cells:
#     if not cell.text.strip() == "":
#         print(cell.text)
# word.save(os.path.join(os.path.dirname(filePath), "mod.docx"))

def messageFile(txtList, msgDir):
    with open(os.path.join(msgDir, "message.txt"), "w", encoding="utf-8") as file:
        file.write("\n".join(txtList))

def docToDocx(filePath):
    import pythoncom
    from win32com import client
    pythoncom.CoInitialize()
    """checking microsoft word was installed"""
    try:
        wrd = client.Dispatch("Word.Application")
    except BaseException as errMsg:
        loggerError.exception("The Word application hasn't been installed!")
        fileDir = os.path.dirname(filePath)
        messageFile(["Не обнаружено установленой программы WORD на вашем компьютере!",
                     str(errMsg)], fileDir)
        return False

    fileDir = os.path.dirname(filePath)
    wrd.visible = 0
    try:
        newFileName = os.path.splitext(filePath)[0]
        doc = wrd.Documents.Open(filePath)
        doc.SaveAs(os.path.join(fileDir, newFileName), FileFormat=12)
        doc.Close()
        isConverted = f"{newFileName}.docx"
    except BaseException as errMsg:
        loggerError.exception("The Word application hasn't been opened or saved!")
        messageFile(["Ошибка открытия или сохранения файла. Закройте все приложения WORD на вашем компьютере!",
                     str(errMsg)], fileDir)
        isConverted = False

    wrd.Quit()
    return isConverted

class WordHandler(FileSystemEventHandler):
    def on_created(self, event):
        """path to file"""
        file = event.src_path
        fileDir = os.path.dirname(file)
        # for file in os.listdir(fileDir):
        """converting *.doc into *.docx"""
        if file.find("~$") == -1:
            """deleting message.txt"""
            if os.path.basename(file) != "message.txt":
                msgPath = os.path.join(fileDir, "message.txt")
                if os.path.exists(msgPath):
                    os.unlink(msgPath)
            if file.endswith(".doc"):
                newFile = docToDocx(os.path.normpath(os.path.join(fileDir, file)))
                if newFile:
                    loggerInfo.info(f"The file {file} was successfully converted to {newFile}")
            elif file.endswith(".docx"):
                splitWordFile(os.path.normpath(os.path.join(fileDir, file)))

class IniHandler(FileSystemEventHandler):
    def __init__(self):
        super().__init__()
        self.obs = None

    def on_modified(self, event):
        if event.src_path.find("settings.ini") != -1:
            config.read("settings.ini")
            obsr = self.obs
            newPath = config.get("Paths", "Path")
            obsr.schedule(WordHandler(), path=os.path.normpath(newPath))
            loggerInfo.info(f'The directory has been changed to {newPath}')

if __name__ == '__main__':
    """directory for observing"""
    observer = Observer()
    observer.schedule(WordHandler(), path=os.path.normpath(config.get("Paths", "Path")))
    observer.start()
    """if settings.ini was changed"""
    observerINI = Observer()
    IniHandler = IniHandler()
    IniHandler.obs = observer
    observerINI.schedule(IniHandler, path=projectDir)
    observerINI.start()

    # try:
    #     while True:
    #         time.sleep(10)
    # except KeyboardInterrupt:
    #     observer.stop()
    #     observerINI.stop()
    #
    observer.join()
    observerINI.join()
