from docx import Document
from xlrd import open_workbook
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import os
from shutil import rmtree as shutil_rmtree
from time import sleep as time_sleep
import logging
from configparser import ConfigParser
from copy import deepcopy

"""win32 service"""
import servicemanager
import socket
import sys
import win32event
import win32service
import win32serviceutil

from concurrent.futures import ProcessPoolExecutor
from multiprocessing import freeze_support

"""global path"""
projectDir = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(
    os.path.abspath(__file__))

"""settings.ini"""
config = ConfigParser()
configList = config.read(os.path.join(projectDir, "settings.ini"))

if configList.__len__() == 0:
    config["DEFAULT"] = {"Path": projectDir}

    with open("settings.ini", "w") as configfile:
        config.write(configfile)

"""Writing the logs"""
loggerError = logging.getLogger("error")
loggerError.setLevel(logging.ERROR)

loggerglobal = logging.getLogger("global")
loggerglobal.setLevel(logging.ERROR)

loggerInfo = logging.getLogger("info")
loggerInfo.setLevel(logging.INFO)

formatter = logging.Formatter("%(asctime)s:%(message)s")

"""if directory Logs doesn't exist"""
logDir = os.path.join(projectDir, "Logs")
if not os.path.isdir(logDir):
    os.mkdir(logDir)

infoHandler = logging.FileHandler(os.path.join(logDir, "info.log"))
infoHandler.setLevel(logging.INFO)
infoHandler.setFormatter(formatter)

errorHandler = logging.FileHandler(os.path.join(logDir, "errors.log"))
errorHandler.setLevel(logging.raiseExceptions)
errorHandler.setFormatter(formatter)

globalHandler = logging.FileHandler(os.path.join(logDir, "global.log"))
globalHandler.setLevel(logging.raiseExceptions)
globalHandler.setFormatter(formatter)

loggerError.addHandler(errorHandler)
loggerInfo.addHandler(infoHandler)
loggerglobal.addHandler(globalHandler)

class valueTable:
    def __init__(self, table):
        self.table = table

    def __getitem__(self, item):
        if isinstance(item, int):
            return self.table[item]
        else:
            fltr = list(filter(lambda x: next(iter(x)) == item, self.table))
            return [next(iter(itm.values())) for itm in fltr]

    def structure(self, mapping=None):
        if isinstance(mapping, list):
            if str(self.table).find("xlrd.sheet.Sheet object") != -1:
                sheet = self.table
                dict_list = []
                for row_index in range(1, sheet.nrows):
                    d = {''.join(sheet.cell(row_index, mapping[0]).value.split()):
                             ''.join(sheet.cell(row_index, mapping[1]).value.split())}
                    dict_list.append(d)
                self.table = dict_list
                return dict_list

class magictree:
    def __init__(self, parent=None):
        self.parent = parent
        self.level = 0 if parent is None else parent.level + 1
        self.attr = []
        self.rows = []

    def add(self, value):
        tr = magictree(self)
        tr.attr.append(value)
        self.rows.append(tr)
        return tr

    def printtree(self):
        def printrows(rows):
            for i in rows:
                print("{}{}".format(i.level * "\t", i.attr))
                printrows(i.rows)

        printrows(self.rows)

"""need an additional class for multiple exceptions"""

class waitexception_1(BaseException):
    pass

class waitexception_3(BaseException):
    pass

def logDecorator(func):
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except BaseException as errMsg:
            loggerglobal.exception(f"An error has been occurred in function {func.__name__}", exc_info=errMsg)

    return wrapper

@logDecorator
def getMappingTable(fileDir):
    pathtofile = os.path.join(fileDir, "mapping.xlsx")
    if not os.path.exists(pathtofile):
        loggerError.error(f"File {pathtofile} not found! Copy a mapping file to the directory!")
        messageFile(["Файл сопоставления оборудования с протоколом не найден!", pathtofile], fileDir)
        return None
    xls = open_workbook(pathtofile)
    sheet = xls.sheet_by_index(0)
    vt = valueTable(sheet)
    vt.structure(mapping=[1, 2])
    return vt

@logDecorator
def replacetext(paragraphs, oldstring='', newstring='', instantreplace=False):
    if instantreplace:
        firstloop = True
        for prg in paragraphs:
            inline = prg.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if firstloop:
                    text = newstring
                    firstloop = False
                else:
                    text = ""
                inline[i].text = text
    else:
        for prg in paragraphs:
            if prg.text.find(oldstring) != -1:
                inline = prg.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if oldstring in inline[i].text:
                        text = inline[i].text.replace(oldstring, newstring)
                        inline[i].text = text

@logDecorator
def mergecells(row, first_merge, last_merge):
    textToDelete = row.cells[last_merge].paragraphs
    mrg = row.cells[first_merge].merge(row.cells[last_merge])
    for prg in mrg.paragraphs:
        for prg_tmp in textToDelete:
            if prg.text.find(prg_tmp.text) != -1:
                paragraphtodelete = prg._element
                paragraphtodelete.getparent().remove(paragraphtodelete)
                paragraphtodelete._p = paragraphtodelete._element = None

    return mrg

@logDecorator
def findparagraph(paragraphs, desc_list, rangeList=None):
    if rangeList is None:
        rangeList = range(0, len(word.paragraphs))

    for desc in desc_list:
        for doc_prg in rangeList:
            conclusion = paragraphs[doc_prg]
            if conclusion.text.lower().find(desc) != -1:
                return conclusion

    return None

@logDecorator
def deleteparagraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

@logDecorator
def replaceparagraph(paragraph, text=''):
    firstloop = True
    inline = paragraph.runs
    # Loop added to work with runs (strings with same style)
    for i in range(len(inline)):
        if firstloop:
            firstloop = False
        else:
            text = ""
        inline[i].text = text

@logDecorator
def get_table_paragrapghs(table, text, row=None, cell=None, get_coordinates=False):
    if isinstance(row, tuple):
        row_range = range(row[0], row[1])
    else:
        row_range = range(0, len(table.rows))

    if isinstance(cell, tuple):
        cell_range = range(cell[0], cell[1])
    else:
        cell_range = range(0, len(table.columns))

    for curr_row in row_range:
        for curr_cell in cell_range:
            curr_value = table.rows[curr_row].cells[curr_cell]
            if curr_value.text.find(text) != -1:
                return (curr_value, (curr_row, curr_cell)) if get_coordinates else curr_value

    return None

@logDecorator
def columns_to_merge(row, text, from_start=True):
    if from_start:
        sequence = range(0, len(row.cells))
    else:
        sequence = range(len(row.cells) - 1, 0, -1)
    for num in sequence:
        if False not in [row.cells[num].text.lower().find(i) != -1 for i in text]:
            return num

    return None

@logDecorator
def rebuildColumns(row=None, row_copy=None, global_vrb=None, isHeader=False):
    merge_start = None
    value_3 = None
    merge_end = None
    value_4_column = None
    value_4 = None
    value_5_column = None

    if isHeader:
        first_col = columns_to_merge(row, ("наименование", "работы"))
        last_col = columns_to_merge(row, ("номер", "протокола"), from_start=False)
        clm_text = ""
        clm_num = 0

        for clm in range(first_col, last_col + 1):
            if clm_text != row.cells[clm].text:
                clm_text = row.cells[clm].text
                clm_num += 1
                if clm_num == 2:
                    merge_start = clm
                    global_vrb.update({"merge_start": clm})
                elif clm_num == 3:
                    value_3 = clm_text
                    global_vrb.update({"value_3_column": clm})
                elif clm_num == 4:
                    merge_end = clm - 1
                    value_4 = clm_text
                    value_4_column = clm
                    global_vrb.update({"merge_end": merge_end})
                    global_vrb.update({"value_4_column": clm})
                elif clm_num == 5:
                    value_5_column = clm
                    global_vrb.update({"value_5_column": clm})
    else:
        merge_start = global_vrb.get("merge_start")
        value_3 = row_copy.cells[global_vrb.get("value_3_column")].text
        merge_end = global_vrb.get("merge_end")
        value_4_column = global_vrb.get("value_4_column")
        value_4 = row_copy.cells[value_4_column].text
        value_5_column = global_vrb.get("value_5_column")

    if not (not merge_start or not merge_end):
        mergecells(row_copy, merge_start, merge_end)
    if not (not value_4_column or not value_3):
        replacetext(row_copy.cells[value_4_column].paragraphs, newstring=value_3, instantreplace=True)
    if not (not value_5_column or not value_4):
        replacetext(row_copy.cells[value_5_column].paragraphs, newstring=value_4, instantreplace=True)

@logDecorator
def splitWordFile(filePath):
    """refreshing the directory
    if directory Logs doesn't exist"""
    fileDir = os.path.dirname(filePath)
    """global variables as dictionary"""
    global global_var
    global_var = {}
    """getting mapping table"""
    mappingTable = getMappingTable(fileDir)
    splitDir = os.path.join(fileDir, os.path.splitext(filePath)[0])
    if os.path.isdir(splitDir):
        for file in os.listdir(splitDir):
            fileToDetete = os.path.join(splitDir, file)
            try:
                if os.path.isdir(fileToDetete):
                    shutil_rmtree(fileToDetete)
                else:
                    os.unlink(fileToDetete)
            except BaseException as errMsg:
                loggerError.exception(f"An error has been occurred deleting the file or the dir: {fileToDetete}")
                messageFile([f"Файл или каталог {fileToDetete} занят другим приложением. Закройте открытые файлы.",
                             str(errMsg)], fileDir)
                return False
    else:
        os.mkdir(splitDir)
    try:
        word = Document(filePath)
    except waitexception_1:
        """trying to wait until OS define the file"""
        time_sleep(1)
        word = Document(filePath)
    except waitexception_2:
        """trying to wait until OS define the file"""
        time_sleep(3)
        word = Document(filePath)
    except BaseException as errMsg:
        loggerError.exception(f"An error occurred while reading file: word = Document({filePath})")
        messageFile(["Ошибка чтения WORD.", str(errMsg), filePath], fileDir)
        return False

    """removing empty rows at the end of file"""
    prg_to_del = []
    for curr_par in range(9, len(word.paragraphs)):
        prg_obj = word.paragraphs[curr_par]
        if not prg_obj.text:
            prg_to_del.append(prg_obj)

    for curr_par in prg_to_del:
        deleteparagraph(curr_par)

    paragraphs = word.tables[0]
    paragraphsCopy = deepcopy(paragraphs)

    """replace the word in the protocol and getting the row of protocol number"""
    protocol_data = get_table_paragrapghs(paragraphsCopy, "Протокол сертификационных испытаний", cell=(0, 5),
                                          get_coordinates=True)
    if protocol_data is not None:
        replacetext(protocol_data[0].paragraphs, "сертификационных ", "")
        global_var.update({"row_protocol_number": protocol_data[1][0]})
    """item name"""
    protocol_name = get_table_paragrapghs(paragraphsCopy, "Тип изделия:", get_coordinates=True)
    if protocol_name is not None:
        global_var.update({"row_protocol_name": paragraphsCopy.rows[protocol_name[1][0]].cells[3].text})

    """deleting paragraphs"""
    if 'small_' in [i.name_val for i in word.styles._element.style_lst]:
        smallstyle = word.styles['small_']
    else:
        smallstyle = word.styles.add_style("small_", word.styles["Normal"].type)
        font = smallstyle.font
        font.size = 88900

    p_1 = findparagraph(word.paragraphs, ["результаты измерений параметров изделий"],
                        range(4, len(word.paragraphs)))
    if p_1 is not None:
        newparagraph = word.add_paragraph(style=smallstyle)
        p_1._element.getparent().replace(p_1._element, newparagraph._element)
    p_2 = findparagraph(word.paragraphs, ["данные протокола могут быть воспроизведены"],
                        range(4, len(word.paragraphs)))
    if p_1 is not None and p_2 is not None:
        deleteparagraph(p_2)
    elif p_1 is None and p_2 is not None:
        newparagraph = word.add_paragraph(style=smallstyle)
        p_2._element.getparent().replace(p_2._element, newparagraph._element)

    """we need this variable to replace the table in the main file"""
    equipmentCopy = deepcopy(word.tables[1])
    rowtodelete = []
    startrow = 0
    rowheader = None
    global newrow
    newrow = None
    secheaderrow = None
    headerCount = 0
    isFirstTable = True
    tree = magictree()
    hierarchy = None
    for row in paragraphs.rows:
        if row.cells[0].text.find("Наименование работы") != -1:
            """getting the count column"""
            count_column = columns_to_merge(row, ("кол-во", "испытанных", "изделий"))
            global_var.update({"count_column": count_column})

            # """merge header"""
            # first_merge = columns_to_merge(row, ("соответствие", "требованиям", "пи"))
            # last_merge = columns_to_merge(row, ("номер", "протокола"), from_start=False)
            # global_var.update({"first_merge": first_merge, "last_merge": last_merge})
            # if not (first_merge is None or last_merge is None):
            #     mergecells(paragraphsCopy.rows[row._index], first_merge, last_merge)
            """moving column's values to the right"""
            rebuildColumns(row=row, row_copy=paragraphsCopy.rows[row._index], global_vrb=global_var, isHeader=True)

            startrow = row._index + 1
            """clearing the paragrapg table"""
            for inx in range(startrow, len(paragraphs.rows)):
                rowtodelete.append(paragraphsCopy.rows[inx]._tr)
            """deleting rows"""
            for delrow in rowtodelete:
                paragraphsCopy._tbl.remove(delrow)
        elif (startrow != 0) and (row._index >= startrow) and (
                row.cells[2].text == row.cells[3].text == row.cells[8].text
                == row.cells[9].text == row.cells[11].text):
            """searching a header if exists"""
            hierarchy = tree.add(paragraphs.rows[row._index])
            """define the header for a conclusion"""
            if (headerCount == 1) or (row.cells[0].text.find("Сертификационные испытания") != -1):
                secheaderrow = row.cells[0].text
            headerCount += 1
        elif (startrow != 0) and (row._index >= startrow):
            if hierarchy is not None:
                hierarchy.add(paragraphs.rows[row._index])
            else:
                tree.add(paragraphs.rows[row._index])

    def outputitems(itemrows):
        global newrow
        for rowlower in itemrows:
            if newrow is None:
                newrow = paragraphsCopy.add_row()
            currentrow = rowlower.attr[0]
            """name for new file(the protocol number"""
            wordname = currentrow.cells[11].text
            """items count"""
            count_column = global_var.get("count_column")
            if count_column is None:
                text_count = ""
            else:
                text_count = currentrow.cells[count_column].text
            """rename protocol string"""
            row_protocol_number = global_var.get("row_protocol_number")
            if row_protocol_number is not None:
                replacetext(paragraphsCopy.rows[row_protocol_number].cells[8].paragraphs,
                            newstring=currentrow.cells[11].text,
                            instantreplace=True)
            """paragraph name"""
            paragraphname = currentrow.cells[0].text
            """moving column's values to the right"""
            rebuildColumns(row_copy=currentrow, global_vrb=global_var)

            newrow._element.getparent().replace(newrow._element, currentrow._element)
            newrow = currentrow

            """replacing the 1-st table"""
            word.tables[0]._element.getparent().replace(word.tables[0]._element, paragraphsCopy._element)

            """if a paragraph isn't found in equipment, deleting the equipment row
            creating a copy of the equipment to edit"""
            if mappingTable is not None:
                equipmenttoedit = deepcopy(equipmentCopy)
                equipmentList = mappingTable[''.join(paragraphname.split())]
                equiptodelete = []
                """need to define do we need a header in the table"""
                equipheader = None
                equipmentrowsexist = False
                for equipmnt in equipmenttoedit.rows:
                    if equipmnt._index > 1:
                        """this is header"""
                        if equipmnt.cells[0].text == equipmnt.cells[1].text:
                            if equipheader is not None and equipmentrowsexist is not True:
                                equiptodelete.append(equipheader)
                            equipmentrowsexist = False
                            equipheader = equipmnt._tr
                        else:
                            if not ''.join(equipmnt.cells[0].text.split()) in equipmentList:
                                equiptodelete.append(equipmnt._tr)
                            elif equipmentrowsexist is False:
                                equipmentrowsexist = True

                """need to check the end of a table"""
                if equipheader is not None and equipmentrowsexist is not True:
                    equiptodelete.append(equipheader)

                for equipdelete in equiptodelete:
                    equipmenttoedit._tbl.remove(equipdelete)

                word.tables[1]._element.getparent().replace(word.tables[1]._element, equipmenttoedit._element)

            """creating conclusion"""
            conclusion = findparagraph(word.paragraphs, ["партия изделий", "выборка в количестве"],
                                       rangeList=range(6, len(word.paragraphs)))

            itemname = global_var.get("row_protocol_name")
            if isFirstTable:
                conclusion_text = f"Партия изделий {itemname} в количестве {text_count} " \
                                  f"шт. прошла входной контроль с положительным результатом."
            else:
                conclusion_text = f"Выборка в количестве {text_count} шт. из партии изделий {itemname} " \
                                  f"прошла сертификационные испытания с положительным результатом."
            replaceparagraph(conclusion, conclusion_text)

            """writing the WORD"""
            word.save(os.path.join(splitDir, f"{wordname}.docx"))

    for rowtree in tree.rows:
        if len(rowtree.rows) > 0:
            """groups"""
            curr_group_raw = rowtree.attr[0]
            """is this a first table"""
            if isFirstTable:
                if curr_group_raw.cells[0].text == secheaderrow:
                    isFirstTable = False

            if rowheader is None:
                rowheader = paragraphsCopy.add_row()
            rowheader._element.getparent().replace(rowheader._element, curr_group_raw._element)
            rowheader = rowtree.attr[0]
            """items"""
            outputitems(rowtree.rows)
        else:
            outputitems(tree.rows)
            break

    return True

@logDecorator
def messageFile(txtList, msgDir):
    with open(os.path.join(msgDir, "message.txt"), "w", encoding="utf-8") as file:
        file.write("\n".join(txtList))

@logDecorator
def docToDocx(filePath):
    import pythoncom
    from win32com import client
    pythoncom.CoInitialize()
    """checking microsoft word was installed"""
    try:
        wrd = client.Dispatch("Word.Application")
    except BaseException as errMsg:
        loggerError.exception("The Word application hasn't been installed!")
        fileDir = os.path.dirname(filePath)
        messageFile(["Не обнаружено установленой программы WORD на вашем компьютере!",
                     str(errMsg)], fileDir)
        return False

    fileDir = os.path.dirname(filePath)
    wrd.visible = 0
    try:
        newFileName = os.path.splitext(filePath)[0]
        doc = wrd.Documents.Open(filePath)
        doc.SaveAs(os.path.join(fileDir, newFileName), FileFormat=12)
        doc.Close()
        isConverted = f"{newFileName}.docx"
    except BaseException as errMsg:
        loggerError.exception("The Word application hasn't been opened or saved!")
        messageFile(["Ошибка открытия или сохранения файла. Закройте все приложения WORD на вашем компьютере!",
                     str(errMsg)], fileDir)
        isConverted = False

    wrd.Quit()
    return isConverted

def getObserveDirectory():
    """if a observe directory doesn't exists, creating"""
    obsDir = os.path.normpath(config.get("DEFAULT", "Path"))
    if not os.path.isdir(obsDir):
        try:
            os.mkdir(obsDir)
        except waitexception_1:
            loggerError.exception(f"An error has been occurred creating the dir: {obsDir}")
            obsDir = os.path.normpath("c:/")

    return obsDir

def do_job(fileDir, file, isDocEnd, isDocxEnd):
    splitCompleted = False
    if isDocEnd:
        newFile = docToDocx(os.path.normpath(os.path.join(fileDir, file)))
        if newFile:
            splitCompleted = splitWordFile(newFile)

    elif isDocxEnd:
        splitCompleted = splitWordFile(os.path.normpath(os.path.join(fileDir, file)))
    if splitCompleted:
        loggerInfo.info("The WORD file has been split successfully.")

class WordHandler(FileSystemEventHandler):
    @logDecorator
    def on_created(self, event):
        """path to file"""
        file = event.src_path
        fileDir = os.path.dirname(file)
        """converting *.doc into *.docx"""
        if file.find("~$") == -1:
            isDocEnd = file.endswith(".doc")
            isDocxEnd = file.endswith(".docx")
            """deleting message.txt"""
            if isDocEnd or isDocxEnd:
                msgPath = os.path.join(fileDir, "message.txt")
                if os.path.exists(msgPath):
                    os.unlink(msgPath)

                # do_job(fileDir=fileDir, file=file, isDocEnd=isDocEnd, isDocxEnd=isDocxEnd)
                with ProcessPoolExecutor() as executor:
                    executor.submit(do_job, fileDir=fileDir, file=file, isDocEnd=isDocEnd, isDocxEnd=isDocxEnd)

class IniHandler(FileSystemEventHandler):
    def __init__(self):
        super().__init__()
        self.obs = None

    @logDecorator
    def on_modified(self, event):
        if event.src_path.find("settings.ini") != -1:
            obsr = self.obs
            config.read(os.path.join(projectDir, "settings.ini"))
            newPath = getObserveDirectory()
            obsr.schedule(WordHandler(), path=newPath)
            loggerInfo.info(f'The directory has been changed to {newPath}')

def obsDirectory(self=None):
    observer = Observer()
    obsDir = getObserveDirectory()
    observer.schedule(WordHandler(), path=obsDir)
    observer.start()
    """if settings.ini was changed"""
    observerINI = Observer()
    IniHandlerVrb = IniHandler()
    IniHandlerVrb.obs = observer
    observerINI.schedule(IniHandlerVrb, path=projectDir)
    observerINI.start()
    try:
        while True:
            if self is not None and self.run_flag is False:
                observer.stop()
                observerINI.stop()
                raise
            time_sleep(10)
    except KeyboardInterrupt:
        observer.stop()
        observerINI.stop()
    observer.join()
    observerINI.join()

class winService(win32serviceutil.ServiceFramework):
    _svc_name_ = "Wordsplit"
    _svc_display_name_ = "Word split"
    _svc_description_ = "Word split application"

    def __init__(self, args):
        win32serviceutil.ServiceFramework.__init__(self, args)
        self.hWaitStop = win32event.CreateEvent(None, 0, 0, None)
        self.run_flag = True
        socket.setdefaulttimeout(60)

    def SvcStop(self):
        self.run_flag = False
        self.ReportServiceStatus(win32service.SERVICE_STOP_PENDING)
        win32event.SetEvent(self.hWaitStop)

    def SvcDoRun(self):
        rc = None
        while rc != win32event.WAIT_OBJECT_0:
            self.main()
            rc = win32event.WaitForSingleObject(self.hWaitStop, 5000)

    def main(self):
        obsDirectory(self)

if __name__ == '__main__':
    freeze_support()
    if len(sys.argv) == 1:
        servicemanager.Initialize()
        servicemanager.PrepareToHostSingle(winService)
        servicemanager.StartServiceCtrlDispatcher()
    else:
        win32serviceutil.HandleCommandLine(winService)
