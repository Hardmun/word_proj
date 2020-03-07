from docx import Document
from os import getcwd
from os import path as os_path
import os
from watchdog.observers import Observer
from watchdog.events import PatternMatchingEventHandler
import sys
import time

def getpath(folder=''):
    return os_path.join(getcwd(), folder)

def splitWordFile():
    word = Document(getpath("files/1.docx"))
    rows = word.tables[0].rows
    # rows[13].delete()

    tbl = word.tables[0]._tbl
    tr = rows[13]._tr
    tbl.remove(tr)

    for row in rows:
        pass
    # cells = word.tables[0]._cells
    # for cell in cells:
    #     if not cell.text.strip() == "":
    #         print(cell.text)
    # word.save(getpath("files/1_.docx"))

# splitWordFile()

# for fl in os.listdir(os.getcwd()):
#     if fl.endswith('.doc'):
#         print(fl)
#         wrd = client.Dispatch("Word.Application")
#         wrd.visible = 0
#         doc = wrd.Documents.Open(full_path("22.doc"))
#         doc.SaveAs(full_path("testconvert"), FileFormat=12)
#         doc.Close()
#         wrd.Quit()

# def projectdir(ctlg, usetempdir=False):
#     """if executable file -  have to change the default path"""
#     if getattr(sys, 'frozen', False) and usetempdir:
#         exe_path = path.dirname(sys.executable)
#         dirPath = path.join(getattr(sys, "_MEIPASS", exe_path), ctlg)
#     else:
#         dirPath = ctlg
#     return dirPath

class MyHandler(PatternMatchingEventHandler):
    patterns = ["*.xml"]

    def process(self, event):
        """
        event.event_type
            'modified' | 'created' | 'moved' | 'deleted'
        event.is_directory
            True | False
        event.src_path
            path/to/observed/file
        """

        with open(event.src_path, 'r') as xml_source:
            pass
            # xml_string = xml_source.read()
            # parsed = xmltodict.parse(xml_string)
            # element = parsed.get('Pulsar', {}).get('OnAir', {}).get('media')
            # if not element:
            #     return

            # media = Media(
            #     title=element.get('title1'),
            #     description=element.get('title3'),
            #     media_id=element.get('media_id1'),
            #     hour=magicdate(element.get('hour')),
            #     length=element.get('title4')
            # )
            # media.save()

    def on_modified(self, event):
        self.process(event)

    def on_created(self, event):
        self.process(event)

if __name__ == '__main__':
    args = sys.argv[1:]
    path_dir=args[0] if args else "."

    observer = Observer()
    observer.schedule(MyHandler(), path=os.path.join("d:/"))
    observer.start()

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()

    observer.join()
